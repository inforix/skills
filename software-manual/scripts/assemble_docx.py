#!/usr/bin/env python3
"""Assemble evidence-backed Markdown sections into an editable Word manual."""

from __future__ import annotations

import argparse
import datetime as dt
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image
except ImportError:  # pragma: no cover - python-docx can still add pictures
    Image = None


SCREENSHOT_RE = re.compile(r'<!--\s*SCREENSHOT:\s*(.*?)\s*-->', re.I)
ATTR_RE = re.compile(r'(\w+)="([^"]*)"')
TABLE_SEP_RE = re.compile(r'^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*$')
SECRET_RE = re.compile(r'(?i)(sk-[a-z0-9_-]{16,}|bearer\s+[a-z0-9._-]{16,}|password\s*[:=]\s*[^\s<]{8,})')
CJK_FONT = "Hiragino Sans GB"
LANG_ZH = True


def tr(zh: str, en: str) -> str:
    return zh if LANG_ZH else en


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--work-dir", required=True, type=Path)
    p.add_argument("--config", type=Path)
    p.add_argument("--sections-dir", type=Path)
    p.add_argument("--screenshots-dir", type=Path)
    p.add_argument("--output", type=Path)
    p.add_argument("--force", action="store_true", help="Allow replacing an existing DOCX")
    return p.parse_args()


def load_json(path: Path, default):
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise SystemExit(f"Cannot read JSON {path}: {exc}") from exc


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.strip())
    run.bold = bold
    run.font.name = CJK_FONT
    rpr = run._element.get_or_add_rPr()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rpr.rFonts.set(qn(f"w:{key}"), CJK_FONT)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "zh-CN")
    rpr.append(lang)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_index, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "E8EEF5")
                cell._tc.get_or_add_tcPr().append(shd)
    if table.rows:
        set_repeat_table_header(table.rows[0])


def column_widths(rows: list[list[str]]) -> list[int]:
    count = max(len(row) for row in rows)
    scores = []
    for idx in range(count):
        values = [row[idx] if idx < len(row) else "" for row in rows]
        scores.append(max(4, min(40, max(len(v) for v in values))))
    total = sum(scores)
    widths = [max(900, round(9360 * score / total)) for score in scores]
    delta = 9360 - sum(widths)
    widths[-1] += delta
    return widths


def set_run_font(run, name="Calibri", size=None, color=None) -> None:
    latin_name = CJK_FONT if name == "Calibri" else name
    run.font.name = latin_name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), latin_name)
    rpr.rFonts.set(qn("w:hAnsi"), latin_name)
    rpr.rFonts.set(qn("w:eastAsia"), CJK_FONT if name in ("Calibri", "Consolas") else name)
    rpr.rFonts.set(qn("w:cs"), latin_name)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "zh-CN")
    rpr.append(lang)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_inline(paragraph, text: str) -> None:
    token_re = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            set_run_font(paragraph.add_run(text[pos:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            set_run_font(run)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, "Consolas", 9)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F2F4F7")
            run._element.get_or_add_rPr().append(shd)
        else:
            label, url = re.match(r'\[([^]]+)\]\(([^)]+)\)', token).groups()
            # Keep readable, offline-safe text; external URL remains visible but nonessential.
            set_run_font(paragraph.add_run(f"{label}（{url}）"))
        pos = match.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]))


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate])
    if placeholder:
        set_run_font(paragraph.add_run(placeholder))
    paragraph.add_run()._r.append(end)


def style_document(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = CJK_FONT
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal._element.rPr.rFonts.set(qn(f"w:{key}"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 30, "1F4D78", 0, 12),
        ("Subtitle", 14, "5B6573", 0, 8),
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = CJK_FONT
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            style._element.rPr.rFonts.set(qn(f"w:{key}"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    code.font.size = Pt(9)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.15)

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_cover(doc: Document, software: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Inches(1.7)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = p.add_run(str(software.get("name", "软件使用手册")))
    title.bold = True
    set_run_font(title, "Calibri", 30, "1F4D78")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = p2.add_run(tr("软件使用手册", "Software User Manual"))
    set_run_font(subtitle, "Calibri", 18, "2E74B5")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version = software.get("version", tr("未标注版本", "Unspecified"))
    desc = software.get("description", "")
    set_run_font(p3.add_run(f"{tr('版本', 'Version')} {version}"), "Calibri", 12, "5B6573")
    if desc:
        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.space_before = Pt(20)
        set_run_font(p4.add_run(str(desc)), "Calibri", 11, "5B6573")
    stamp = doc.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stamp.paragraph_format.space_before = Inches(1.4)
    set_run_font(stamp.add_run(dt.date.today().isoformat()), "Calibri", 10, "7A8491")


def configure_body_section(section, software: dict) -> None:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(hp.add_run(f"{software.get('name', '软件手册')} · {software.get('version', '')}"), "Calibri", 9, "5B6573")
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(fp, " PAGE ", "1")
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:start"), "1")
    section._sectPr.append(pg_num_type)


def add_code_block(doc: Document, code_text: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F5F7")
    p._p.get_or_add_pPr().append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "D9DEE5")
        borders.append(border)
    p._p.get_or_add_pPr().append(borders)
    for idx, line in enumerate(code_text.splitlines() or [""]):
        if idx:
            p.add_run().add_break()
        safe = "\u200b".join(re.findall(r'.{1,90}', line)) if len(line) > 90 else line
        set_run_font(p.add_run(safe), "Consolas", 9)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    configure_table(table, [9360])
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    add_inline(p, text)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EEF5FB")
    cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def resolve_screenshot(marker: dict, screenshot_dir: Path, manifest: dict) -> Path | None:
    entry = manifest.get(marker.get("id", ""), {})
    filename = entry.get("file") or marker.get("file") or (marker.get("id", "") + ".png")
    candidate = screenshot_dir / filename
    if entry.get("status", "captured") != "captured" or not candidate.is_file():
        return None
    return candidate


def image_width_inches(path: Path) -> float:
    if Image is None:
        return 6.1
    try:
        with Image.open(path) as image:
            dpi_info = image.info.get("dpi", (96, 96))
            dpi_x = dpi_info[0] or 96
            dpi_y = (dpi_info[1] if len(dpi_info) > 1 else dpi_x) or 96
            natural_width = image.width / dpi_x
            aspect_width_for_max_height = 7.0 * image.width / image.height * (dpi_y / dpi_x)
            return min(6.1, natural_width, max(1.0, aspect_width_for_max_height))
    except Exception:
        return 6.1


def add_screenshot(doc: Document, marker: dict, path: Path | None, figure_no: int, warnings: list[str]) -> None:
    description = marker.get("description") or marker.get("alt") or marker.get("id") or "界面截图"
    if path is None:
        add_callout(doc, tr(f"待补截图：{description}（未作为已验证截图计入手册）", f"Screenshot required: {description} (not counted as verified)"))
        warnings.append(f"Missing screenshot: {marker.get('id', description)}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(image_width_inches(path)))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", marker.get("id", description))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_after = Pt(8)
    cap = caption.add_run(f"{tr('图', 'Figure')} {figure_no}  {description}")
    cap.italic = True
    set_run_font(cap, "Calibri", 9, "5B6573")


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_markdown(doc: Document, text: str, screenshot_dir: Path, manifest: dict, stats: dict, warnings: list[str]) -> None:
    lines = text.splitlines()
    idx = 0
    in_code = False
    code_lines: list[str] = []
    while idx < len(lines):
        line = lines[idx]
        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                stats["code_blocks"] += 1
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        marker_match = SCREENSHOT_RE.search(line)
        if marker_match:
            marker = dict(ATTR_RE.findall(marker_match.group(1)))
            path = resolve_screenshot(marker, screenshot_dir, manifest)
            if path:
                stats["screenshots"] += 1
            stats["figure_markers"] += 1
            add_screenshot(doc, marker, path, stats["figure_markers"], warnings)
            idx += 1
            continue

        if line.startswith("|") and idx + 1 < len(lines) and TABLE_SEP_RE.match(lines[idx + 1]):
            block = [line]
            idx += 2
            while idx < len(lines) and lines[idx].startswith("|"):
                block.append(lines[idx])
                idx += 1
            rows = parse_table(block)
            cols = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx in range(cols):
                    set_cell_text(table.cell(r_idx, c_idx), row[c_idx] if c_idx < len(row) else "", r_idx == 0)
            configure_table(table, column_widths(rows))
            stats["tables"] += 1
            continue

        heading = re.match(r'^(#{1,6})\s+(.+?)\s*#*$', line)
        if heading:
            level = min(3, len(heading.group(1)))
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading.group(2))
            stats["headings"] += 1
            idx += 1
            continue
        bullet = re.match(r'^\s*[-*+]\s+(.+)$', line)
        number = re.match(r'^\s*\d+[.)]\s+(.+)$', line)
        if bullet or number:
            p = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            add_inline(p, (bullet or number).group(1))
            idx += 1
            continue
        if line.startswith(">"):
            add_callout(doc, line.lstrip("> "))
            idx += 1
            continue
        if not line.strip():
            idx += 1
            continue
        p = doc.add_paragraph()
        add_inline(p, line.strip())
        idx += 1
    if in_code:
        add_code_block(doc, "\n".join(code_lines))
        stats["code_blocks"] += 1


def main() -> int:
    global CJK_FONT, LANG_ZH
    args = parse_args()
    work_dir = args.work_dir.resolve()
    config_path = (args.config or work_dir / "manual-config.json").resolve()
    config = load_json(config_path, None)
    if not isinstance(config, dict):
        raise SystemExit(f"Missing or invalid config: {config_path}")
    sections_dir = (args.sections_dir or work_dir / "sections").resolve()
    screenshots_dir = (args.screenshots_dir or work_dir / "screenshots").resolve()
    output_cfg = config.get("output", {})
    CJK_FONT = str(output_cfg.get("cjk_font") or CJK_FONT)
    LANG_ZH = str(config.get("software", {}).get("language", "zh-CN")).lower().startswith("zh")
    filename = output_cfg.get("filename", "software-manual.docx")
    if not str(filename).lower().endswith(".docx"):
        filename = Path(str(filename)).with_suffix(".docx").name
    output = (args.output or work_dir / filename).resolve()
    if output.exists() and not args.force:
        raise SystemExit(f"Output exists; pass --force to replace: {output}")
    output.parent.mkdir(parents=True, exist_ok=True)

    manifest_data = load_json(screenshots_dir / "screenshots-manifest.json", {"screenshots": []})
    manifest = {str(item.get("id")): item for item in manifest_data.get("screenshots", []) if isinstance(item, dict)}
    sections = config.get("sections") or []
    if not sections:
        sections = [{"file": path.name, "title": path.stem, "start_on_new_page": True} for path in sorted(sections_dir.glob("*.md"))]
    missing = [str(sections_dir / item["file"]) for item in sections if not (sections_dir / item.get("file", "")).is_file()]
    if missing:
        raise SystemExit("Missing section files:\n" + "\n".join(missing))

    combined = "\n".join((sections_dir / item["file"]).read_text(encoding="utf-8") for item in sections)
    if SECRET_RE.search(combined):
        raise SystemExit("Potential secret detected in source sections; redact before assembly")

    doc = Document()
    style_document(doc)
    software = config.get("software", {})
    doc.core_properties.title = f"{software.get('name', tr('软件', 'Software'))} {tr('使用手册', 'User Manual')}"
    doc.core_properties.subject = tr("软件使用手册", "Software User Manual")
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "software manual, user guide, 软件手册"
    add_cover(doc, software)
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_body_section(body_section, software)

    if output_cfg.get("include_toc", True):
        doc.add_heading(tr("目录", "Contents"), level=1)
        for item in sections:
            entry = doc.add_paragraph()
            entry.paragraph_format.left_indent = Inches(0.2)
            entry.paragraph_format.space_after = Pt(4)
            add_inline(entry, str(item.get("title") or Path(item.get("file", "Section")).stem))
        toc = doc.add_paragraph()
        add_field(toc, ' TOC \\o "1-3" \\h \\z \\u ')
        doc.add_page_break()

    stats = {"sections": 0, "headings": 0, "tables": 0, "screenshots": 0, "figure_markers": 0, "code_blocks": 0}
    warnings: list[str] = []
    for section_index, item in enumerate(sections):
        source = sections_dir / item["file"]
        text = source.read_text(encoding="utf-8")
        if not text.strip():
            warnings.append(f"Empty section: {source.name}")
            continue
        if section_index and item.get("start_on_new_page", True):
            doc.add_page_break()
        if not re.match(r'^#\s+', text.lstrip()):
            doc.add_heading(str(item.get("title") or source.stem), level=1)
            stats["headings"] += 1
        add_markdown(doc, text, screenshots_dir, manifest, stats, warnings)
        stats["sections"] += 1

    doc.save(output)
    size_mb = output.stat().st_size / (1024 * 1024)
    max_mb = float(output_cfg.get("max_size_mb", 50))
    if size_mb > max_mb:
        warnings.append(f"DOCX size {size_mb:.2f} MB exceeds configured {max_mb:.2f} MB")
    report = {
        "status": "completed_with_warnings" if warnings else "completed",
        "primary_output_file": str(output),
        "primary_format": "docx",
        **stats,
        "size_bytes": output.stat().st_size,
        "warnings": warnings,
    }
    (work_dir / "docx-build-report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
