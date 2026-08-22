#!/usr/bin/env python3
"""Validate structural, accessibility, and offline-safety properties of a DOCX manual."""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
}
W = "{" + NS["w"] + "}"
PR = "{" + NS["pr"] + "}"
PLACEHOLDER_RE = re.compile(r'(?i)(\bTODO\b|\bTBD\b|\[\[TOC\]\]|lorem ipsum|待替换|your[_ -]?token)')
SECRET_RE = re.compile(r'(?i)(sk-[a-z0-9_-]{16,}|bearer\s+[a-z0-9._-]{16,}|password\s*[:=]\s*[^\s<]{8,})')
DANGEROUS_FIELD_RE = re.compile(r'(?i)(INCLUDETEXT|INCLUDEPICTURE|DDEAUTO?|MACROBUTTON)')


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--docx", required=True, type=Path)
    p.add_argument("--config", type=Path)
    p.add_argument("--report", type=Path)
    return p.parse_args()


def add_issue(items: list[dict], code: str, message: str) -> None:
    items.append({"code": code, "message": message})


def xml_root(zf: zipfile.ZipFile, name: str):
    return ET.fromstring(zf.read(name))


def text_content(root) -> str:
    return "".join(node.text or "" for node in root.iter(W + "t"))


def validate(args: argparse.Namespace) -> dict:
    errors: list[dict] = []
    warnings: list[dict] = []
    docx_path = args.docx.resolve()
    if not docx_path.is_file():
        add_issue(errors, "missing_file", f"DOCX not found: {docx_path}")
        return {"status": "failed", "errors": errors, "warnings": warnings}

    try:
        document = Document(docx_path)
    except Exception as exc:
        add_issue(errors, "python_docx_open_failed", str(exc))
        return {"status": "failed", "errors": errors, "warnings": warnings}

    required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml", "word/_rels/document.xml.rels"}
    try:
        with zipfile.ZipFile(docx_path) as zf:
            names = set(zf.namelist())
            missing = sorted(required - names)
            if missing:
                add_issue(errors, "missing_ooxml_parts", ", ".join(missing))
            forbidden = [name for name in names if re.search(r'(vbaProject\.bin|word/embeddings/|activeX/)', name, re.I)]
            if forbidden:
                add_issue(errors, "active_content", ", ".join(forbidden))

            root = xml_root(zf, "word/document.xml")
            all_text = text_content(root)
            if PLACEHOLDER_RE.search(all_text):
                add_issue(errors, "placeholder", "Unresolved placeholder text detected")
            if SECRET_RE.search(all_text):
                add_issue(errors, "secret", "Potential credential detected")

            field_text = " ".join(node.text or "" for node in root.iter(W + "instrText"))
            if "TOC" not in field_text:
                add_issue(errors, "missing_toc", "TOC field is missing")
            if "PAGE" not in field_text and not any(name.startswith("word/footer") for name in names):
                add_issue(errors, "missing_page_number", "PAGE field is missing")
            if DANGEROUS_FIELD_RE.search(field_text):
                add_issue(errors, "dangerous_field", "Dangerous Word field detected")

            headings = []
            list_items = 0
            tables = 0
            for para in root.iter(W + "p"):
                ppr = para.find(W + "pPr")
                if ppr is None:
                    continue
                style = ppr.find(W + "pStyle")
                if style is not None:
                    val = style.get(W + "val", "")
                    if val.lower().startswith("heading"):
                        headings.append(val)
                    if val.lower().startswith("list"):
                        list_items += 1
                if ppr.find(W + "numPr") is not None:
                    list_items += 1
            if not headings:
                add_issue(errors, "missing_headings", "No real Heading styles found")
            for tbl in root.iter(W + "tbl"):
                tables += 1
                pr = tbl.find(W + "tblPr")
                grid = tbl.find(W + "tblGrid")
                if pr is None or pr.find(W + "tblW") is None or pr.find(W + "tblInd") is None or grid is None:
                    add_issue(errors, "table_geometry", "A table lacks explicit tblW/tblInd/tblGrid")
                    break
                for cell in tbl.iter(W + "tc"):
                    tc_pr = cell.find(W + "tcPr")
                    if tc_pr is None or tc_pr.find(W + "tcW") is None:
                        add_issue(errors, "cell_geometry", "A table cell lacks explicit tcW")
                        break

            drawings = list(root.iter("{" + NS["wp"] + "}inline"))
            anchors = list(root.iter("{" + NS["wp"] + "}anchor"))
            if anchors:
                add_issue(errors, "floating_images", f"Found {len(anchors)} floating image(s); use inline images")
            missing_alt = 0
            for doc_pr in root.iter("{" + NS["wp"] + "}docPr"):
                if not (doc_pr.get("descr") or "").strip():
                    missing_alt += 1
            if missing_alt:
                add_issue(errors, "missing_alt_text", f"{missing_alt} image(s) lack alt text")

            for rel_name in [name for name in names if name.endswith(".rels")]:
                rel_root = xml_root(zf, rel_name)
                for rel in rel_root.iter(PR + "Relationship"):
                    rel_type = rel.get("Type", "")
                    target_mode = rel.get("TargetMode", "")
                    rel_type_lower = rel_type.lower()
                    if "attachedtemplate" in rel_type_lower or "oleobject" in rel_type_lower or rel_type_lower.endswith("/package"):
                        add_issue(errors, "dangerous_relationship", f"{rel_name}: {rel_type}")
                    if target_mode == "External" and "hyperlink" not in rel_type.lower():
                        add_issue(errors, "external_dependency", f"{rel_name}: {rel.get('Target', '')}")
    except (zipfile.BadZipFile, KeyError, ET.ParseError) as exc:
        add_issue(errors, "invalid_ooxml", str(exc))
        headings, list_items, tables, drawings = [], 0, 0, []

    if args.config:
        try:
            config = json.loads(args.config.read_text(encoding="utf-8"))
            expected = [item.get("title", "") for item in config.get("sections", [])]
            rendered_text = "\n".join(p.text for p in document.paragraphs)
            for title in expected:
                if title and title not in rendered_text:
                    add_issue(warnings, "configured_title_not_literal", f"Configured title not found literally: {title}")
        except (OSError, json.JSONDecodeError) as exc:
            add_issue(errors, "config_read_failed", str(exc))

    if list_items == 0 and re.search(r'(?m)^\s*(?:[-*•]|\d+[.)])\s+', "\n".join(p.text for p in document.paragraphs)):
        add_issue(errors, "fake_lists", "List-looking text exists without Word numbering")

    status = "failed" if errors else ("passed_with_warnings" if warnings else "passed")
    return {
        "status": status,
        "primary_output_file": str(docx_path),
        "primary_format": "docx",
        "structural_validation": "failed" if errors else "passed",
        "render_validation": "required_separately",
        "page_visual_inspection": "required_separately",
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "headings": len(headings),
        "numbered_or_bulleted_paragraphs": list_items,
        "inline_images": len(drawings),
        "errors": errors,
        "warnings": warnings,
    }


def main() -> int:
    args = parse_args()
    report = validate(args)
    report_path = (args.report or args.docx.with_name("docx-validation-report.json")).resolve()
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 1 if report["status"] == "failed" else 0


if __name__ == "__main__":
    sys.exit(main())
