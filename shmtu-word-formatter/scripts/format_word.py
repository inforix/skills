#!/usr/bin/env python3
"""按上海海事大学党政公文格式刷 Word。"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

TITLE_FONT = "方正小标宋简体"
BODY_FONT = "仿宋_GB2312"
LEVEL1_FONT = "黑体"
LEVEL2_FONT = "楷体_GB2312"
LEVEL34_FONT = BODY_FONT
PAGE_FONT = "宋体"

TITLE_SIZE = Pt(22)  # 2号
BODY_SIZE = Pt(16)  # 3号
PAGE_SIZE = Pt(14)  # 4号

LEVEL1_PATTERN = re.compile(r"^[一二三四五六七八九十百千]+、")
LEVEL2_PATTERN = re.compile(r"^（[一二三四五六七八九十百千]+）")
LEVEL3_PATTERN = re.compile(r"^\d+\.")
LEVEL4_PATTERN = re.compile(r"^（\d+）")

DATE_PATTERN = re.compile(r"(\d{4})年0?(\d{1,2})月0?(\d{1,2})日")


def normalize_date_text(text: str) -> str:
    def repl(match: re.Match[str]) -> str:
        year, month, day = match.group(1), int(match.group(2)), int(match.group(3))
        return f"{year}年{month}月{day}日"

    return DATE_PATTERN.sub(repl, text)


def get_or_create_rfonts(run) -> OxmlElement:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    return r_fonts


def set_run_font(run, font_name: str, size: Pt, bold: bool = False) -> None:
    run.font.size = size
    run.bold = bold
    run.font.name = font_name
    r_fonts = get_or_create_rfonts(run)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def clear_paragraph_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)


def apply_text_style(paragraph, font_name: str, size: Pt, bold: bool = False) -> None:
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_font(run, font_name, size, bold=bold)


def classify_font_for_paragraph(text: str) -> tuple[str, bool]:
    stripped = text.strip()
    if LEVEL1_PATTERN.match(stripped):
        return LEVEL1_FONT, False
    if LEVEL2_PATTERN.match(stripped):
        return LEVEL2_FONT, False
    if LEVEL3_PATTERN.match(stripped) or LEVEL4_PATTERN.match(stripped):
        return LEVEL34_FONT, False
    return BODY_FONT, False


def set_page_layout(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, PAGE_FONT, PAGE_SIZE)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)


def set_footer_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        clear_paragraph_runs(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        left = paragraph.add_run("— ")
        set_run_font(left, PAGE_FONT, PAGE_SIZE)
        add_page_field(paragraph)
        right = paragraph.add_run(" —")
        set_run_font(right, PAGE_FONT, PAGE_SIZE)


def build_document_from_text(raw_text: str) -> Document:
    doc = Document()
    lines = raw_text.splitlines()

    if not lines:
        lines = [""]

    for line in lines:
        doc.add_paragraph(line)

    return doc


def first_non_empty_paragraph_index(doc: Document) -> int:
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip():
            return idx
    doc.add_paragraph("")
    return len(doc.paragraphs) - 1


def apply_gov_style(doc: Document, title_override: str | None, add_page_numbers: bool) -> None:
    set_page_layout(doc)

    # 日期格式归一化
    for p in doc.paragraphs:
        normalized = normalize_date_text(p.text)
        if normalized != p.text:
            p.text = normalized

    title_idx = first_non_empty_paragraph_index(doc)
    title_paragraph = doc.paragraphs[title_idx]
    if title_override:
        title_paragraph.text = title_override.strip()

    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.first_line_indent = Pt(0)
    apply_text_style(title_paragraph, TITLE_FONT, TITLE_SIZE)

    for idx, p in enumerate(doc.paragraphs):
        if idx == title_idx:
            continue

        text = p.text.strip()
        if not text:
            continue

        p.paragraph_format.first_line_indent = Pt(32)  # 左空二字（约）
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        font_name, bold = classify_font_for_paragraph(text)
        apply_text_style(p, font_name, BODY_SIZE, bold=bold)

    if add_page_numbers:
        set_footer_page_numbers(doc)


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError("unknown", b"", 0, 1, f"无法解码文本文件: {path}")


def load_document(input_path: Path | None, inline_text: str | None, text_file: Path | None) -> Document:
    if inline_text:
        return build_document_from_text(inline_text)

    if text_file:
        return build_document_from_text(read_text_file(text_file))

    if not input_path:
        raise ValueError("请提供 --input（docx/txt/md）或 --text / --text-file")

    suffix = input_path.suffix.lower()
    if suffix == ".docx":
        return Document(str(input_path))

    return build_document_from_text(read_text_file(input_path))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="按上海海事大学党政公文格式刷 Word")
    parser.add_argument("--input", type=Path, help="输入文件（.docx/.txt/.md）")
    parser.add_argument("--output", type=Path, required=True, help="输出 .docx 文件")
    parser.add_argument("--text", type=str, help="直接传入文本内容")
    parser.add_argument("--text-file", type=Path, help="从文本文件读取内容")
    parser.add_argument("--title", type=str, help="可选：覆盖标题")
    parser.add_argument("--no-page-number", action="store_true", help="不写入页码")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    doc = load_document(args.input, args.text, args.text_file)

    apply_gov_style(
        doc,
        title_override=args.title,
        add_page_numbers=not args.no_page_number,
    )

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))
    print(f"已生成: {args.output}")


if __name__ == "__main__":
    main()
